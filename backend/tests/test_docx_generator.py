from datetime import date

from docx import Document

from app.models.department import Department
from app.models.education_history import EducationHistory
from app.models.employee import Employee
from app.models.enums import EducationLevel, EmploymentStatus, Gender, RelativeType
from app.models.position import Position
from app.models.relative import Relative
from app.models.work_history import WorkHistory
from app.services.docx_generator import generate_objektivka


def _make_employee(**overrides) -> Employee:
    defaults = {
        "last_name": "Aliyev",
        "first_name": "O'ktam",
        "middle_name": "Bahodir o'g'li",
        "birth_date": date(1990, 5, 12),
        "birth_place": "Toshkent",
        "gender": Gender.MALE,
        "employment_status": EmploymentStatus.ACTIVE,
    }
    defaults.update(overrides)
    return Employee(**defaults)


def test_generates_valid_docx_zip_container() -> None:
    buf = generate_objektivka(_make_employee())
    data = buf.read()
    assert data[:2] == b"PK"
    assert len(data) > 0


def test_script_tag_content_is_not_leaked_into_document() -> None:
    employee = _make_employee(
        positive_traits="<p>Masuliyatli</p><script>alert(document.cookie)</script>",
    )
    doc = Document(generate_objektivka(employee))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "alert" not in full_text
    assert "document.cookie" not in full_text
    assert "Masuliyatli" in full_text


def test_paragraph_boundaries_are_preserved() -> None:
    employee = _make_employee(positive_traits="<p>Birinchi jumla</p><p>Ikkinchi jumla</p>")
    doc = Document(generate_objektivka(employee))
    paragraph_texts = [p.text for p in doc.paragraphs if p.text.strip()]

    assert "Birinchi jumla" in paragraph_texts
    assert "Ikkinchi jumla" in paragraph_texts
    assert "Birinchi jumlaIkkinchi jumla" not in paragraph_texts


def test_matches_reference_template_title_and_layout() -> None:
    """Tashkilotning haqiqiy МАЪЛУМОТНОМА namunasiga mos: kirill sarlavha, F.I.Sh.
    qatori, va juft-ustunli (label\\tlabel / value\\tvalue) shaxsiy ma'lumotlar."""
    employee = _make_employee()
    doc = Document(generate_objektivka(employee))
    texts = [p.text for p in doc.paragraphs]

    assert texts[0] == "МАЪЛУМОТНОМА"
    assert texts[1] == employee.full_name
    assert "Туғилган йили:\tТуғилган жойи:" in texts
    assert "12.05.1990\tToshkent" in texts


def test_timeline_merges_education_and_work_history_sorted_by_date() -> None:
    """"МЕҲНАТ ФАОЛИЯТИ" — haqiqiy namunadagidek ta'lim va mehnat tarixi ikkita alohida
    ro'yxat emas, bitta sana bo'yicha saralangan xronologik ro'yxatga birlashadi."""
    employee = _make_employee()
    employee.work_history = [
        WorkHistory(
            organization_name="Kadr MChJ",
            position_title="Dasturchi",
            start_date=date(2018, 9, 1),
            end_date=None,
        )
    ]
    employee.education_history = [
        EducationHistory(
            institution_name="TATU",
            specialty="Kompyuter injiniringi",
            level=EducationLevel.BACHELOR,
            start_date=date(2013, 9, 1),
            end_date=date(2017, 6, 30),
        )
    ]

    doc = Document(generate_objektivka(employee))
    texts = [p.text for p in doc.paragraphs]
    timeline_index = texts.index("МЕҲНАТ ФАОЛИЯТИ")

    assert texts[timeline_index + 1] == "2013-2017 йй. - TATU, Kompyuter injiniringi йўналиши"
    assert texts[timeline_index + 2] == "2018-ҳ.в. йй. - Kadr MChJ, Dasturchi"


def test_relatives_table_matches_reference_columns() -> None:
    employee = _make_employee()
    employee.relatives = [
        Relative(
            relation_type=RelativeType.FATHER,
            full_name="Aliyev Bahodir",
            birth_year=1965,
            birth_place="Namangan",
            workplace="Maktab",
            position_title="O'qituvchi",
            address="Namangan",
        )
    ]

    doc = Document(generate_objektivka(employee))
    table = doc.tables[0]

    assert [c.text for c in table.rows[0].cells] == [
        "Қариндош-\nлари",
        "Фамилияси, исми ва\nотасининг исми",
        "Туғилган йили\nва жойи",
        "Иш жойи ва\nлавозими",
        "Яшаш жойи",
    ]
    assert [c.text for c in table.rows[1].cells] == [
        "Отаси",
        "Aliyev Bahodir",
        "1965 йил, Namangan",
        "Maktab, O'qituvchi",
        "Namangan",
    ]


def test_current_position_line_uses_since_date_and_department() -> None:
    employee = _make_employee(position_since=date(2021, 9, 6))
    employee.position = Position(id=1, title="Bosh mutaxassis", department_id=1)
    employee.position.department = Department(id=1, name="Kadrlar bo'limi")

    doc = Document(generate_objektivka(employee))
    texts = [p.text for p in doc.paragraphs]

    assert "2021 йил 6-Сентябрдан:" in texts
    assert "Kadrlar bo'limi Bosh mutaxassis" in texts
