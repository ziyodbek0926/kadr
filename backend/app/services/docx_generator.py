"""Xodim ma'lumotlari asosida "МАЪЛУМОТНОМА" (obyektivka, .docx) hujjatini python-docx
bilan generatsiya qilish. Tuzilma va maydonlar tashkilotning haqiqiy namunaviy hujjatidan
(kirill yozuvida) olingan: sarlavha, F.I.Sh. qatori, joriy lavozim + "...дан:" sanasi,
juft-ustunli shaxsiy ma'lumotlar, birlashtirilgan ta'lim+mehnat xronologiyasi ("МЕҲНАТ
ФАОЛИЯТИ") va yaqin qarindoshlar jadvali — aynan shu tartibda va ustun nomlari bilan.

Employee obyekti chaqirilishdan oldin barcha bog'liq to'plamlari (relatives/
education_history/work_history/awards/...) yuklangan bo'lishi kerak
(app.crud.employee.CRUDEmployee.get_detail orqali)."""

import re
from datetime import date
from io import BytesIO

import nh3
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Cm, Inches, Pt

from app.models.education_history import EducationHistory
from app.models.employee import Employee
from app.models.relative import Relative
from app.models.work_history import WorkHistory

_EMPTY = "Йўқ"
_TAB_POSITION = Inches(3.3)

_RELATION_LABELS = {
    "father": "Отаси",
    "mother": "Онаси",
    "spouse": "Турмуш ўртоғи",
    "child": "Фарзанди",
    "sibling": "Ака/укаси, опа/синглиси",
    "other": "Қариндоши",
}
_EDUCATION_LEVEL_LABELS = {
    "secondary": "ўрта",
    "secondary_special": "ўрта-махсус",
    "bachelor": "олий",
    "master": "олий",
    "phd": "олий (илмий даража)",
}
# "...дан:" (ablative — "shu sanadan beri") qo'shimchasi bilan, haqiqiy namunadagi
# "6-Сентябрдан:" uslubiga mos keladigan shaklda.
_MONTH_SINCE_CY = {
    1: "январдан", 2: "февралдан", 3: "мартдан", 4: "апрелдан",
    5: "майдан", 6: "июндан", 7: "июлдан", 8: "августдан",
    9: "сентябрдан", 10: "октябрдан", 11: "ноябрдан", 12: "декабрдан",
}


def _strip_html(value: str | None) -> list[str]:
    """Rich-text maydonini oddiy paragraflarga aylantiradi — python-docx HTML'ni o'zi
    render qila olmaydi. Qiymat odatda app.utils.sanitize orqali saqlashdan oldin
    tozalangan bo'ladi, lekin bu yerda YANA nh3 orqali o'tkaziladi (ikkinchi himoya
    qatlami): oddiy regex bilan teg olib tashlansa, `<script>alert(1)</script>` kabi
    holatlarda teg o'chadi-yu, ICHIDAGI matn ("alert(1)") qolib ketadi — nh3 esa
    script/style kabi xavfli elementlarni mazmuni bilan birga olib tashlaydi (faqat
    "p"/"br" ruxsat etilgan holda ham)."""
    if not value:
        return []
    safe_html = nh3.clean(value, tags={"p", "br"})
    text = re.sub(r"<br\s*/?>", "\n", safe_html)
    text = re.sub(r"</p\s*>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    return [line.strip() for line in text.split("\n") if line.strip()]


def _format_date(value: date | None) -> str:
    return value.strftime("%d.%m.%Y") if value else _EMPTY


def _set_base_style(document: DocxDocument) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1)


def _add_centered(document: DocxDocument, text: str, size: int, bold: bool = True) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _add_field_pair(
    document: DocxDocument,
    label1: str,
    value1: str | None,
    label2: str | None = None,
    value2: str | None = None,
) -> None:
    """Haqiqiy namunadagi uslub: bir qatorda qalin yozuvli 1-2 ta bo'lim nomi, undan
    keyingi qatorda TAB orqali tekislangan qiymatlar. Table emas — original hujjat ham
    aynan shu tarzda (TAB belgilari bilan) qurilgan."""
    label_p = document.add_paragraph()
    label_p.paragraph_format.tab_stops.add_tab_stop(_TAB_POSITION, WD_TAB_ALIGNMENT.LEFT)
    label_run = label_p.add_run(f"{label1}:")
    label_run.bold = True
    if label2:
        label_p.add_run("\t")
        label_run2 = label_p.add_run(f"{label2}:")
        label_run2.bold = True

    value_p = document.add_paragraph()
    value_p.paragraph_format.tab_stops.add_tab_stop(_TAB_POSITION, WD_TAB_ALIGNMENT.LEFT)
    value_p.paragraph_format.space_after = Pt(6)
    value_p.add_run(value1 or _EMPTY)
    if label2:
        value_p.add_run("\t")
        value_p.add_run(value2 or _EMPTY)


def _education_end_date(edu: EducationHistory) -> date:
    # Faqat end_date is not None bilan oldindan filtrlangan ro'yxatlarda chaqiriladi —
    # shu sababli bu yerda None kelmasligi kafolatlangan (mypy uchun aniq tur beradi).
    assert edu.end_date is not None
    return edu.end_date


def _since_line(position_since: date | None) -> str | None:
    if not position_since:
        return None
    month = _MONTH_SINCE_CY[position_since.month].capitalize()
    return f"{position_since.year} йил {position_since.day}-{month}:"


def _build_timeline(
    education_history: list[EducationHistory], work_history: list[WorkHistory]
) -> list[str]:
    """"МЕҲНАТ ФАОЛИЯТИ" bo'limi — haqiqiy namunada ta'lim va mehnat tarixi ALOHIDA
    jadvallarda emas, balki bitta xronologik ro'yxatda (sana bo'yicha saralangan)
    birlashtirilgan holda ko'rsatiladi."""
    entries: list[tuple[date, date | None, str]] = []
    for edu in education_history:
        if not edu.start_date:
            continue
        description = edu.institution_name
        if edu.specialty:
            description += f", {edu.specialty} йўналиши"
        entries.append((edu.start_date, edu.end_date, description))
    for work in work_history:
        entries.append((work.start_date, work.end_date, f"{work.organization_name}, {work.position_title}"))

    entries.sort(key=lambda entry: entry[0])
    lines = []
    for start, end, description in entries:
        end_label = str(end.year) if end else "ҳ.в."
        lines.append(f"{start.year}-{end_label} йй. - {description}")
    return lines


def _relative_row(relative: Relative) -> list[str]:
    birth = ""
    if relative.birth_year:
        birth = f"{relative.birth_year} йил"
        if relative.birth_place:
            birth += f", {relative.birth_place}"
    elif relative.birth_place:
        birth = relative.birth_place

    workplace = ", ".join(part for part in (relative.workplace, relative.position_title) if part)
    relation_value = relative.relation_type.value
    return [
        _RELATION_LABELS.get(relation_value, relation_value),
        relative.full_name,
        birth or _EMPTY,
        workplace or _EMPTY,
        relative.address or _EMPTY,
    ]


def _add_table(document: DocxDocument, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def generate_objektivka(employee: Employee) -> BytesIO:
    document = Document()
    _set_base_style(document)

    _add_centered(document, "МАЪЛУМОТНОМА", 16)
    _add_centered(document, employee.full_name, 16)

    if employee.position:
        since = _since_line(employee.position_since)
        if since:
            since_p = document.add_paragraph()
            since_run = since_p.add_run(since)
            since_run.bold = True
        document.add_paragraph(f"{employee.position.department.name} {employee.position.title}")

    document.add_paragraph()

    _add_field_pair(document, "Туғилган йили", _format_date(employee.birth_date), "Туғилган жойи", employee.birth_place)
    _add_field_pair(document, "Миллати", employee.nationality, "Партиявийлиги", employee.party_affiliation)

    completed_education = [edu for edu in employee.education_history if edu.end_date is not None]
    latest_education = max(completed_education, key=_education_end_date, default=None)
    level_label = _EDUCATION_LEVEL_LABELS.get(latest_education.level.value, "") if latest_education else None
    completed = (
        f"{_education_end_date(latest_education).year} йил {latest_education.institution_name}"
        if latest_education
        else None
    )
    _add_field_pair(document, "Маълумоти", level_label, "Тамомлаган", completed)

    _add_field_pair(document, "Илмий даражаси", employee.academic_degree, "Илмий унвони", employee.academic_title)
    _add_field_pair(
        document,
        "Қайси чет тилларини билади",
        employee.foreign_languages,
        "Ҳарбий (махсус) унвон",
        employee.military_rank,
    )

    awards_summary = ", ".join(award.name for award in employee.awards) if employee.awards else None
    _add_field_pair(document, "Давлат мукофотлари билан тақдирланганми (қанақа)", awards_summary)
    _add_field_pair(
        document,
        "Халқ депутатлари, республика, вилоят, шаҳар ва туман Кенгашлари депутатлари "
        "ёки бошқа сайланадиган органлар аъзосими",
        employee.public_office_note,
    )

    document.add_paragraph()
    _add_centered(document, "МЕҲНАТ ФАОЛИЯТИ", 14)
    timeline = _build_timeline(employee.education_history, employee.work_history)
    if timeline:
        for line in timeline:
            document.add_paragraph(line)
    else:
        document.add_paragraph("Маълумот киритилмаган.")

    document.add_paragraph()
    document.add_paragraph()
    _add_centered(document, f"{employee.full_name}нинг яқин қариндошлари тўғрисида", 12)
    _add_centered(document, "МАЪЛУМОТ", 12)
    document.add_paragraph()

    if employee.relatives:
        _add_table(
            document,
            [
                "Қариндош-\nлари",
                "Фамилияси, исми ва\nотасининг исми",
                "Туғилган йили\nва жойи",
                "Иш жойи ва\nлавозими",
                "Яшаш жойи",
            ],
            [_relative_row(relative) for relative in employee.relatives],
        )
    else:
        document.add_paragraph("Маълумот киритилмаган.")

    positive_lines = _strip_html(employee.positive_traits)
    if positive_lines:
        document.add_paragraph()
        _add_centered(document, "Ижобий фазилатлари", 12)
        for line in positive_lines:
            document.add_paragraph(line)

    negative_lines = _strip_html(employee.negative_traits)
    if negative_lines:
        document.add_paragraph()
        _add_centered(document, "Салбий фазилатлари", 12)
        for line in negative_lines:
            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
